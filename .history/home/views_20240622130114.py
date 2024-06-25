from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, redirect
from .models import FormularioUser
import pandas as pd
import shopify
import unidecode
from docx import Document
import requests
import json
from django.conf import settings
from django.views.decorators.cache import cache_page
from django.views.decorators.csrf import csrf_exempt
from django.urls import reverse
import random
import string
import threading  # Importe o módulo threading
from django.contrib.auth.models import User
from django.contrib.auth import authenticate
from django.contrib.auth import login as login_django
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache

############################
 #  WEBHOOK KIWIFY
############################
@csrf_exempt
def webhook_handler(request):
    if request.method == 'POST':
        # Extrair dados do corpo da solicitação POST (JSON)
        payload = json.loads(request.body)

        # Extrair nome e email do payload (assumindo que estão presentes nos dados recebidos)
        customer = payload.get('Customer', {})
        nome = customer.get('full_name', '')
        email = customer.get('email', '')

        if nome and email:
            # Gerar nome de usuário
            username = (nome[:3] + ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))).lower()

            # Gerar senha aleatória de 8 caracteres alfanuméricos
            senha = ''.join(random.choices(string.ascii_letters + string.digits, k=8))

            # Criar um novo usuário no Django User
            user = User.objects.create_user(username=username, email=email, password=senha)

            # Criar ou atualizar o FormularioUser
            formulario_user, created = FormularioUser.objects.update_or_create(
                email=email,
                defaults={
                    'nome': nome,
                    'username': username,
                    'senha': senha  # Salvar a senha não criptografada no modelo FormularioUser
                }
            )

            # Responder ao webhook com uma mensagem de sucesso
            return JsonResponse({
                'status': 'success',
                'message': 'Usuário criado com sucesso!',
                'username': username,
                'senha': senha
            }, status=200)

    # Se a solicitação não for POST ou os dados estiverem incompletos, retornar erro
    return JsonResponse({'status': 'error', 'message': 'Dados inválidos'}, status=400)
###########################
                            ###  FAZER LOGIN #####
###########################
def login(request):
    if request.method == 'GET':
        return render(request, 'login.html')
    else:
        username = request.POST.get('username')
        senha = request.POST.get('senha')
        user = authenticate(username=username, password=senha)
        if user:
            login_django(request, user)
            
            # Buscar ou criar o FormularioUser
            formulario_user, created = FormularioUser.objects.get_or_create(email=user.email, defaults={'nome': user.get_full_name()})
            if created:
                formulario_user.save()
                print("Novo FormularioUser criado para o usuário.")
            else:
                # Verificar se a loja já foi produzida
                if formulario_user.lojaproduzida:
                    return render(request, 'modal5.html')
                elif formulario_user.banners:
                    return render(request, 'produzir.html')
                elif formulario_user.produt:
                    return render(request, 'tema_cor.html')
                elif formulario_user.plan:
                    return render(request, 'validacao.html')
            return redirect('pais')
        else:
            return render(request, 'modal.html')

@login_required
def redirecionar_para_login(request):
    return redirect('login')

###########################
                            ### SELECIONAR PAÍS ######
###########################
@login_required
def pais(request): 
    return render(request, 'pais.html')
@login_required
def sele_pais(request):
    if request.method == "POST":
        pais = request.POST.get('pais')
        print("Pais escolhido:", pais)
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            
            if formulario_user:
                formulario_user.pais = pais
                formulario_user.save()
                print("Formulário atualizado com o país escolhido.")
                return redirect('indice')
            else:
                print("Formulário não encontrado.")
                return redirect('login')                
        else:
            return redirect('login')
@login_required
def redirecionar_para_pais(request):
    return redirect('pais')

############################
                            ### INDICE CRIAR CONTA E OBTER URL ######
############################
@login_required
def indice(request):
    if request.method == 'GET': 
        return render(request, 'indice.html')
    else:
        link_store = request.POST.get('link_store')
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            
            if formulario_user:
                if link_store and link_store.startswith("https://admin.shopify.com/store/"):
                    loja_id = link_store.split("https://admin.shopify.com/store/")[1]
                    url_loja = f"{loja_id}.myshopify.com"
                    formulario_user.link_store = link_store
                    formulario_user.url_loja = url_loja
                    formulario_user.save()
                    print("Formulário atualizado, link e url")
                    if request.user.is_authenticated:
                        email = request.user.email
                        formulario_user = FormularioUser.objects.filter(email=email).first()
                        pais = formulario_user.pais
                        if pais == "Brasil":
                            return redirect('nicho')
                        else:
                            return redirect('acessos')
                    else:
                        print("Formulário não encontrado.") 
                        return redirect('indice')
                else:
                    print("URL inválida")
                    return render(request, 'indice.html', {'error': 'A URL deve começar com "https://admin.shopify.com/store/"'})
            else:
                print("Formulário não encontrado.") 
                return redirect('indice')
        else:
            return redirect('login')
@login_required
def redirecionar_para_indice(request):    
    return redirect('indice')
############################
                            ### SELECIONAR NICHO ######
############################
@login_required
def nicho(request): 
    return render(request, 'nicho.html')
@login_required
def redirecionar_para_nicho(request):
    return redirect('nicho')
@login_required
def sele_nicho(request):
    if request.method == "POST":
        nicho = request.POST.get('nicho')
        print("Valor de nicho:", nicho)  # Imprime o valor de 'nicho' no console
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            if formulario_user:
                formulario_user.nicho = nicho
                formulario_user.save()
                print("Formulário atualizado com o nicho escolhido.")
            else:
                print("Formulário não encontrado.") 
            # Redirecionar para a página de validação ou outra página desejada
            return redirect('acessos')  
        else:
            return redirect('login') 
############################
                            ### CRIAR APP PRIVADO E VERIFICAR ######
############################
@login_required
def verificando(request):    
    if request.method == 'POST': 
        link_store = request.POST.get('link_store')
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()

            link_store = formulario_user.link_store
            appprivado = f"{link_store}/settings/apps/development"
            return redirect(appprivado) 
@login_required
def redirecionar_para_verificando(request):
    return redirect('verificando')
def verificar_shopify(SHOP_URL, API_KEY, PRIVATE_APP_PASSWORD):
    endpoint = f"https://{SHOP_URL}/admin/api/2024-04/shop.json"
    headers = {'Content-Type': 'application/json'}
    auth = (API_KEY, PRIVATE_APP_PASSWORD)

    try:
        response = requests.get(endpoint, headers=headers, auth=auth)
        if response.status_code == 200:
            return True
        else:
            return False
    except requests.exceptions.RequestException as e:
        print(f"Erro ao conectar à API da Shopify: {e}")
        return False
@login_required
def acessos(request):
    if request.method == "GET":
        return render(request, 'acessos.html')
    else:
        token_senha = request.POST.get('token_senha')
        chave_de_api = request.POST.get('chave_de_api')
        chave_secreta = request.POST.get('chave_secreta')    

        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            
            if formulario_user:
                SHOP_URL = formulario_user.url_loja

                # Verificar se as credenciais da Shopify são válidas
                if verificar_shopify(SHOP_URL, chave_de_api, token_senha):
                    formulario_user.token_senha = token_senha
                    formulario_user.chave_de_api = chave_de_api
                    formulario_user.chave_secreta = chave_secreta
                    formulario_user.save()
                    print("Formulário atualizado com tokens de acesso.")
                    if request.user.is_authenticated:
                        email = request.user.email
                        formulario_user = FormularioUser.objects.filter(email=email).first()
                        if formulario_user:
                            # Marcar a loja como produzida
                            formulario_user.plan = True
                            formulario_user.save()
                            return redirect('validacao')  # Exemplo de redirecionamento após selecionar o nicho
                else:
                    print("Credenciais inválidas.")
                    return render(request, 'model6.html')
            else:
                print("Formulário não encontrado.") 
                return redirect('login')
        else:
            return redirect('login')
@login_required
def redirecionar_para_acessos(request):
    return redirect('acessos')   
############################
                            ### ACESSAR E VALIDAR CONTA SHOPIFY ######
############################
@login_required
def validacao(request):
    if request.method == 'GET': 
        return render(request, 'validacao.html')
    else:
        link_store = request.POST.get('link_store')
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()

            link_store = formulario_user.link_store
            validar = f"{link_store}/settings/plan"
            return redirect(validar) 
        else:
            print("Formulário não encontrado.") 
            return redirect('login')  
@login_required
def redirecionar_para_validacao(request):
    return redirect('validacao')
############################
                            ### ADICIONAR PRODUTOS ######
############################
def importar_produtos(df, SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD, API_KEY, API_SECRET):
    shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
    session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
    shopify.ShopifyResource.activate_session(session)

    for index, row in df.iterrows():
        handle = row['Handle']
        title = row['Title']
        image_src = row['Image Src']
        
        if pd.isna(title) or title.strip() == '':
            title = 'Título padrão'

        product = shopify.Product({
            "title": title,
            "body_html": row['Body (HTML)'] if pd.notna(row['Body (HTML)']) else '',
            "vendor": row['Vendor'] if pd.notna(row['Vendor']) else '',
            "product_type": row['Type'] if pd.notna(row['Type']) else '',
            "tags": unidecode.unidecode(row['Tags']).lower().replace(' ', '-') + ",mais-vendidos" if pd.notna(
                row['Tags']) else '',
            "published": row['Published'] if pd.notna(row['Published']) else True,
        })
        product.save()

        # Adicionar a variante ao produto
        variant = shopify.Variant({
            "taxable": 'false',
            "price": str(row['Variant Price']) if pd.notna(row['Variant Price']) else '0.00',
            "compare_at_price": str(row['Variant Compare At Price']) if pd.notna(
                row['Variant Compare At Price']) else '0.00',
            "inventory_quantity": 1000,
            "inventory_management": 'shopify',
            "inventory_policy": 'continue'
        })
        product.variants = [variant]
        product.save()

        # Adicionar a imagem ao produto
        if pd.notna(image_src):
            image = shopify.Image({'src': image_src})
            image.product_id = product.id
            image.save()

    shopify.ShopifyResource.clear_session()
def importar_produtos_async(df, SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD, API_KEY, API_SECRET):
    # Inicia a importação de produtos em uma thread separada
    thread = threading.Thread(target=importar_produtos, args=(df, SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD, API_KEY, API_SECRET))
    thread.start()
@login_required
def ver_conta(request):
    if request.method == "POST":        
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            estilo = formulario_user.nicho
            pais = formulario_user.pais
            SHOP_URL = formulario_user.url_loja
            PRIVATE_APP_PASSWORD = formulario_user.token_senha
            API_KEY = formulario_user.chave_de_api
            API_SECRET = formulario_user.chave_secreta                        
            API_VERSION = '2024-04'

            # Definir o caminho do arquivo CSV com base no nicho
            if pais == "Brasil":
                if estilo == "Genérica":
                    df_path = "static/home/csv/30GENERICOS.csv"
                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/30ELETRONICOS.csv"
                elif estilo == "Kids":
                    df_path = "static/home/csv/30KIDS.csv"
                elif estilo == "Casa":
                    df_path = "static/home/csv/30CASA.csv"
                elif estilo == "Pet":
                    df_path = "static/home/csv/30PETS.csv"
                elif estilo == "Fitness":
                    df_path = "static/home/csv/30FITNESS.csv"
                elif estilo == "Masculino":
                    df_path = "static/home/csv/30MASCULINO.csv"
                elif estilo == "Feminino":
                    df_path = "static/home/csv/30FEMININO.csv"
            elif pais == "Portugal":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Paquistão":
                df_path = "static/home/csv/30GEN_PAKISTAN.csv"
            elif pais == "Índia":
                df_path = "static/home/csv/30GEN_INDIA.csv"
            elif pais == "Argentina":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Bolívia":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Chile":
                ddf_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Colômbia":
                df_path = "static/home/csv/30GENERICOScol.csv"
            elif pais == "Equador":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Paraguai":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Peru":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Uruguai":
                df_path = "static/home/csv/30GENERICOS.csv"
            elif pais == "Venezuela":
                df_path = "static/home/csv/30GENERICOS.csv"

            df = pd.read_csv(df_path)

            # Verificar informações da loja na Shopify
            endpoint = f"https://{SHOP_URL}/admin/api/{API_VERSION}/shop.json"
            headers = {'Content-Type': 'application/json'}
            auth = (API_KEY, PRIVATE_APP_PASSWORD)
            response = requests.get(endpoint, headers=headers, auth=auth)

            if response.status_code == 200:
                shop_info = response.json()
                plan_name = shop_info.get('shop', {}).get('plan_name')

                if not plan_name:
                    return redirect('modal2')  # Plan name não encontrado

                if plan_name in ['affiliate', 'basic', 'Shopify', 'Advanced']:
                    importar_produtos_async(df, SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD, API_KEY, API_SECRET)
                    if request.user.is_authenticated:
                        email = request.user.email
                        formulario_user = FormularioUser.objects.filter(email=email).first()
                        if formulario_user:
                            # Marcar a loja como produzida
                            formulario_user.produt = True
                            formulario_user.save()
                            return redirect('prefer')                    
                elif plan_name == 'starter':
                    return redirect('modal4')  
                elif plan_name == 'trial':
                    return redirect('pagina_erro')
                else:
                    return redirect('pagina_erro')
            else:
                return redirect('modal2') 
        else:
            print("Formulário não encontrado.") 
            return redirect('login')
    else:
        return render(request, 'acessos.html')
############################
                            ### PREFERENCIAS DESATIVAR SENHA ######
############################
def prefer(request):
    if request.method == 'GET': 
        return render(request, 'prefer.html')
    else:
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            link_store = formulario_user.link_store
            linkpref = f"{link_store}/online_store/preferences"
            return redirect(linkpref)
        else:
            print("Formulário não encontrado.") 
            return redirect('login') 
def redirecionar_para_prefer(request):    
    return redirect('prefer')
############################                            
                            ### ESCOLHER COR SUBIR TEMA PAGINAS ######       
############################
@login_required
def aguarde(request):
    return render(request, 'tema_cor.html')
@login_required
def redirecionar_para_aguarde(request):
    return redirect('aguarde')
@login_required
def tema(request):  
    if request.method == "POST":
        cor = request.POST.get('cor')
        print("Cor escolhida:", cor)
        
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            
            if formulario_user:
                formulario_user.cor = cor
                formulario_user.save()
                print("Formulário atualizado com a cor escolhida.")
                
                # Recuperar as outras informações necessárias
                estilo = formulario_user.nicho
                pais = formulario_user.pais
                paleta = formulario_user.cor
                SHOP_URL = formulario_user.url_loja
                PRIVATE_APP_PASSWORD = formulario_user.token_senha
                API_KEY = formulario_user.chave_de_api
                API_SECRET = formulario_user.chave_secreta
                API_VERSION = '2024-04'
       
                ##############################################################
                if pais == "Brasil":
                    Tema = "https://github.com/ronildob/TEMAS-SHOPIFY/raw/master/tema_padr%C3%A3o.zip"
                    paginas = {
                            "Rastrear Pedidos": "static/home/politicas/Rastrear Pedido.txt",
                            "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                            "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                            "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                            "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                            "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                            "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                            "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                        }                                    
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                        "Terms of service": "static/home/politicas/Termos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                        "Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                        "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                    }
                    if estilo == "Genérica":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa6.png?raw=true"                   
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa6.png?raw=true" 
                    
                    elif estilo == "Eletrônicos":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/ROXO/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_ELETRONICOS/blob/main/ELETR%C3%94NICOS/VERMELHO/CAPAS/capa6.png?raw=true" 
                            
                    elif estilo == "Kids":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20DESKTOP/desktop3.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20DESKTOP/desktop4.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/ROXO/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_KIDS/blob/main/KIDS/VERMELHO/CAPAS/capa6.png?raw=true" 
                        
                    elif estilo == "Casa":
                        if paleta == "Paleta de Cor 1":  # LARANJA 
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/ROXO/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/BANNER%20MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_CASA/blob/main/CASA/VERMELHO/CAPAS/capa6.png?raw=true" 
                    
                    elif estilo == "Pet":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/ROXO/CAPAS/capa6.png?raw=true"                              
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa4_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/Nicho_Pets/blob/main/PETS/VERMELHO/CAPAS/capa6.png?raw=true"  
                    
                    elif estilo == "Fitness":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERDE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/ROXO/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FITNESS/blob/main/FITNESS/VERMELHO/CAPAS/capa6.png?raw=true" 
                    
                    elif estilo == "Masculino":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/LARANJA/CAPAS/capa6.png?raw=true"                    
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERDE/MOBILE/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/ROXO/CAPAS/capa6.png?raw=true"                 
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/PRETO/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_MASCULINO/blob/main/NICHO%20MASCULINO/VERMELHO/CAPAS/capa6.png?raw=true" 
                    
                    elif estilo == "Feminino":
                        if paleta == "Paleta de Cor 1":  # LARANJA
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/LARANJA/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true"                  
                        elif paleta == "Paleta de Cor 2": # VERDE
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERDE/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true"
                        elif paleta == "Paleta de Cor 3": # ROXO
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/ROXO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true"                
                        elif paleta == "Paleta de Cor 4": # PRETO
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/PRETO/MOBILE/mobile2.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true"
                        elif paleta == "Paleta de Cor 5": # AZUL
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true" 
                        elif paleta == "Paleta de Cor 6": # VERMELHO
                            desktop1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/DESKTOP/desktop1.png?raw=true"
                            desktop2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/DESKTOP/desktop2.png?raw=true"
                            desktop3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/DESKTOP/desktop3.png?raw=true"
                            mobile1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/MOBILE/mobile1.png?raw=true"
                            mobile2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/MOBILE/mobile2.png?raw=true"
                            mobile3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/VERMELHO/MOBILE/mobile3.png?raw=true"
                            capa1_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa1.png?raw=true"
                            capa2_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa2.png?raw=true"
                            capa3_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa3.png?raw=true"
                            capa4_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa4.png?raw=true"
                            capa5_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa5.png?raw=true"
                            capa6_url = "https://github.com/ronildob/NICHO_FEMININO/blob/main/NICHO%20FEMININO/AZUL/CAPAS/capa6.png?raw=true"          
                    ##########################                                         
                elif pais == "Portugal":
                    Tema = "https://github.com/ronildob/TEMAS-SHOPIFY/raw/master/tema_padr%C3%A3o.zip"
                    paginas = {
                            "Rastrear Pedidos": "static/home/politicas/Rastrear Pedido.txt",
                            "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                            "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                            "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                            "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                            "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                            "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                            "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                        }                                    
                    politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                    "Terms of service": "static/home/politicas/Termos de Uso.docx",
                    "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                    "Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                    "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                    } 
                    if paleta == "Paleta de Cor 1":  # LARANJA
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/BANNER%20MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa6.png?raw=true"                   
                    elif paleta == "Paleta de Cor 2": # VERDE
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 3": # ROXO
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa6.png?raw=true"                 
                    elif paleta == "Paleta de Cor 4": # PRETO
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/BANNER%20MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 5": # AZUL
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Mobile/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 6": # VERMELHO
                        desktop1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa6.png?raw=true"                        
                elif pais == "Paquistão" or pais == "Índia":
                    Tema = "https://github.com/ronildob/tema_ingles/raw/main/THEME_INGLES.zip"
                    paginas = {
                            "Track Orders": "static/home/politicas/Track Orders.txt",
                            "E-Commerce Agreement": "static/home/politicas/E-Commerce Agreement.docx",
                            "Terms of Use": "static/home/politicas/Terms of Use.docx",
                            "Deadline": "static/home/politicas/Deadline.docx",
                            "Privacy Policy": "static/home/politicas/Privacy Policy.docx",
                            "Exchanges or Return": "static/home/politicas/Exchanges or Return.docx",
                            "About Us": "static/home/politicas/About Us.docx",
                            "Secure Payment": "static/home/politicas/Secure Payment.docx",
                        }                                    
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/E-Commerce Agreement.docx",
                        "Terms of service": "static/home/politicas/Terms of Use.docx",
                        "Shipping policy": "static/home/politicas/Deadline.docx",
                        "Privacy policy": "static/home/politicas/Privacy Policy.docx",
                        "Refund policy": "static/home/politicas/Exchanges or Return.docx",
                    }                                         
                    if paleta == "Paleta de Cor 1":  # LARANJA
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/desktop/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/desktop/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/desktop/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/LARANJA/mobile/mobile3.png?raw=true" 
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa6.png?raw=true"                   
                    elif paleta == "Paleta de Cor 2": # VERDE
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/desktop/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/desktop/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/desktop/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERDE/mobile/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 3": # ROXO
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/desktop/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/desktop/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/desktop/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/ROXO/mobile/mobile3.png?raw=true"            
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa6.png?raw=true"                 
                    elif paleta == "Paleta de Cor 4": # PRETO
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/desktop/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/desktop/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/desktop/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/PRETO/mobile/mobile3.png?raw=true"            
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 5": # AZUL
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/desktop/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/desktop/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/desktop/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/mobile/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/mobile/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/AZUL/mobile/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 6": # VERMELHO
                        desktop1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banners-ingles-/blob/main/BANNERS%20INGL%C3%8AS/GEN%C3%89RICA/VERMELHO/BANNER%20MOBILE/mobile3.png?raw=true"            
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa6.png?raw=true" 
                else:
                    Tema = "https://github.com/ronildob/TEMAS-SHOPIFY/raw/master/tema-padrao-espanhol.zip"
                    paginas = {
                            "Rastrear Pedidos": "static/home/politicas/Rastrear Pedido.txt",
                            "Contrato de Comercio Electrónico": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                            "Términos de Uso": "static/home/politicas/Términos de Uso.docx",
                            "Plazo de Entrega": "static/home/politicas/Plazo de Entrega.docx",
                            "Política de Privacidad": "static/home/politicas/Política de Privacidad.docx",
                            "Cambios o Devoluciones": "static/home/politicas/Cambios o Devoluciones.docx",
                            "Sobre Nosotros": "static/home/politicas/Sobre Nosotros.docx",
                            "Pago Seguro": "static/home/politicas/Pago Seguro.docx",
                        }                                    
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                        "Terms of service": "static/home/politicas/Términos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Plazo de Entrega.docx",
                        "Privacy policy": "static/home/politicas/Política de Privacidad.docx",
                        "Refund policy": "static/home/politicas/Cambios o Devoluciones.docx",
                    }                                         
                    if paleta == "Paleta de Cor 1":  # LARANJA
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/LARANJA/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/LARANJA/CAPAS/capa6.png?raw=true"                   
                    elif paleta == "Paleta de Cor 2": # VERDE
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERDE/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERDE/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 3": # ROXO
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/ROXO/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/ROXO/CAPAS/CAPAS/capa6.png?raw=true"                 
                    elif paleta == "Paleta de Cor 4": # PRETO
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/PRETO/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/PRETO/CAPAS/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 5": # AZUL
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/AZUL/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/AZUL/Capas/capa6.png?raw=true" 
                    elif paleta == "Paleta de Cor 6": # VERMELHO
                        desktop1_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/DESKTOP/desktop1.png?raw=true"
                        desktop2_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/DESKTOP/desktop2.png?raw=true"
                        desktop3_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/DESKTOP/desktop3.png?raw=true"
                        mobile1_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/MOBILE/mobile1.png?raw=true"
                        mobile2_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/MOBILE/mobile2.png?raw=true"
                        mobile3_url = "https://github.com/ronildob/banner_espanhol/blob/main/VERMELHO/MOBILE/mobile3.png?raw=true"
                        capa1_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa1.png?raw=true"
                        capa2_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa2.png?raw=true"
                        capa3_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa3.png?raw=true"
                        capa4_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa4.png?raw=true"
                        capa5_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa5.png?raw=true"
                        capa6_url = "https://github.com/ronildob/NICHO_GENERICA/blob/main/GEN%C3%89RICA/VERMELHO/CAPAS/capa6.png?raw=true"
                #############################
                #############################    
                shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
                session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
                shopify.ShopifyResource.activate_session(session)
                #############################
                def importar_tema(tema):
                    new_theme = shopify.Theme()
                    new_theme.name = estilo
                    new_theme.src = Tema
                    new_theme.role = "main"
                    if new_theme.save():
                        print("Tema criado com sucesso. ID:", new_theme.id)
                    else:
                        print("Falha ao criar o tema. Erros:", new_theme.errors.full_messages())               
                def subir_banners(urls):
                    for nome_arquivo, url in urls.items():
                        mutation = """
                                    mutation fileCreate($files: [FileCreateInput!]!) {
                                        fileCreate(files: $files) {
                                        files {
                                            id
                                            alt
                                        }
                                        userErrors {
                                            message
                                        }
                                        }
                                    }
                                    """

                        file_data = {
                            "files": [
                                {
                                    "alt": f"{nome_arquivo} para shopify",
                                    "originalSource": url
                                }
                            ]
                        }

                        response = requests.post(f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",
                                                json={"query": mutation, "variables": file_data},
                                                headers={"X-Shopify-Access-Token": PRIVATE_APP_PASSWORD})

                        if response.status_code == 200:
                            data = response.json()
                            if "errors" in data:
                                print(f"Erro ao criar arquivo: {data['errors']}")
                            else:
                                file_info = data["data"]["fileCreate"]["files"][0]
                                link = f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json?query={file_info['id']}"
                                print(f"Arquivo {nome_arquivo} criado com sucesso! Link: {link}")
                        else:
                            print(f"Erro ao enviar a solicitação para criar arquivo: {response.status_code}")
                def criar_paginas(paginas):
                    for titulo, caminho_arquivo in paginas.items():
                        if caminho_arquivo.endswith('.docx'):
                            doc = Document(caminho_arquivo)
                            conteudo = '\n'.join([p.text for p in doc.paragraphs])
                        elif caminho_arquivo.endswith('.txt'):
                            with open(caminho_arquivo, 'r') as f:
                                conteudo = f.read()
                        else:
                            print(f"Formato de arquivo não suportado: {caminho_arquivo}")
                            continue
                        nova_pagina = shopify.Page()
                        nova_pagina.title = titulo
                        nova_pagina.body_html = conteudo
                        nova_pagina.save()
                    pass
                def ler_documento(caminho):
                    document = Document(caminho)
                    texto = ""
                    for paragraph in document.paragraphs:
                        texto += paragraph.text + "\n"
                    return texto
                def atualizar_politicas():
                    # Endpoint para atualizar as políticas
                    endpoint = f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json"

                    # Cabeçalhos da requisição com autenticação
                    headers = {
                        'X-Shopify-Access-Token': PRIVATE_APP_PASSWORD,
                        'Content-Type': 'application/json'
                    }

                    # Mutation para atualizar a política
                    mutation = """
                            mutation shopPolicyUpdate($shopPolicy: ShopPolicyInput!) {
                            shopPolicyUpdate(shopPolicy: $shopPolicy) {
                                shopPolicy {
                                id
                                }
                                userErrors {
                                field
                                message
                                }
                            }
                            }
                            """

                    # Variáveis da mutação
                    for titulo, caminho_documento in politicas_para_atualizar.items():
                        corpo = ler_documento(caminho_documento)
                        politica_para_atualizar = {
                            "body": corpo,
                            "type": titulo.upper().replace(" ", "_")
                        }

                        variables = {
                            "shopPolicy": politica_para_atualizar
                        }

                        # Envia a requisição POST com a mutação
                        response = requests.post(endpoint, json={"query": mutation, "variables": variables},
                                                headers=headers)

                        # Verifica se a requisição foi bem sucedida
                        if response.status_code == 200:
                            data = response.json()
                            if 'errors' in data:
                                print(f"Erro ao atualizar a política '{titulo}':")
                                for error in data['errors']:
                                    print(f"Erro: {error['message']}")
                            else:
                                print(f"Política '{titulo}' atualizada com sucesso!")
                        else:
                            print(
                                f"Erro ao enviar a requisição para a política '{titulo}': Status Code - {response.status_code}, Response Text - {response.text}")
                #############################
                importar_tema(Tema)
                urls = {
                    "desktop1.png": desktop1_url,
                    "desktop2.png": desktop2_url,
                    "desktop3.png": desktop3_url,
                    "mobile1.png": mobile1_url,
                    "mobile2.png": mobile2_url,
                    "mobile3.png": mobile3_url,
                    "capa1.png": capa1_url,
                    "capa2.png": capa2_url,
                    "capa3.png": capa3_url,
                    "capa4.png": capa4_url,
                    "capa5.png": capa5_url,
                    "capa6.png": capa6_url
                }
                subir_banners(urls)
                criar_paginas(paginas)
                atualizar_politicas()
                shopify.ShopifyResource.clear_session()
                if request.user.is_authenticated:
                    email = request.user.email
                    formulario_user = FormularioUser.objects.filter(email=email).first()
                    if formulario_user:
                        # Marcar a loja como produzida
                        formulario_user.banners = True
                        formulario_user.save()
                        if pais == "Brasil":                    
                            return redirect('dsers')
                        else:
                            return redirect('produzir')

        else:
            print("Formulário não encontrado.") 
            return redirect('login')
############################
                             ### INSTALAR DSERS ###### 
############################
@login_required
def dsers(request):
    return render(request, 'dsers.html')
@login_required
def redirecionar_para_dsers(request):    
    return redirect('dsers')


@login_required
def redirecionar_para_dsers(request):
    return redirect('dsers')
############################
                            ### CUSTOMIZAR TEMA E COLEÇÕES ###### 
############################

@login_required
@never_cache
def produzir(request): 
    if request.method == "GET":
        return render(request, 'produzir.html')
    else:
        telefone = request.POST.get('telefone')
        email_suporte = request.POST.get('email_suporte')
        empresa = request.POST.get('empresa')
        business_hours = request.POST.get('business_hours')    
        print("dados do coletados:", telefone, email_suporte, empresa, business_hours) 
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            if formulario_user:
                formulario_user.telefone = telefone
                formulario_user.email_suporte = email_suporte
                formulario_user.empresa = empresa
                formulario_user.business_hours = business_hours
                formulario_user.save()
                print("Formulário atualizado com dados do usuario.")
            try: 
                SHOP_URL = formulario_user.url_loja
                PRIVATE_APP_PASSWORD = formulario_user.token_senha
                API_KEY = formulario_user.chave_de_api
                API_SECRET = formulario_user.chave_secreta 
                estilo = formulario_user.nicho
                pais = formulario_user.pais
                paleta = formulario_user.cor 
                telefone = formulario_user.telefone
                email_suporte = formulario_user.email_suporte
                empresa = formulario_user.empresa
                business_hours = formulario_user.business_hours
                whatsapp = str(telefone)                     
                API_VERSION = '2024-04'
                cod_pais = {
                    "Argentina": "+54",
                    "Bolívia": "+591",
                    "Brasil": "+55",
                    "Chile": "+56",
                    "Colômbia": "+57",
                    "Equador": "+593",
                    "Paraguai": "+595",
                    "Peru": "+51",
                    "Uruguai": "+598",
                    "Venezuela": "+58",
                    "Paquistão": "+92",
                    "Índia": "+91",
                    "Portugal": "+351"
                }
                if pais in cod_pais:
                    tel_whats = cod_pais[pais] + str(telefone)
                    whatsapp_completo = cod_pais[pais] + whatsapp
                if pais == "Brasil" or pais == "Portugal":
                    if business_hours == hora1:
                        antend = "Segunda a Sexta das 9:00 as 18:00"
                    elif business_hours == hora2: 
                        antend = "Segunda a Sexta das 9:00 as 22:00"
                    elif business_hours == hora3:
                        antend = "Segunda a Sábado das 9:00 as 18:00"
                    elif business_hours == hora4:
                        antend = "Segunda a Sábado das 9:00 as 22:00"
                    elif business_hours == hora5:
                        antend = "Segunda a Sexta 9:00 as 18:00, Sábado das 9:00 as 12:00"
                    elif business_hours == hora6:
                        antend = "Segunda a Sexta 9:00 as 22:00, Sábado das 9:00 as 12:00"
                    elif business_hours == hora7:
                        antend = "Todos os dias das 9:00 as 18:00"
                    elif business_hours == hora8:
                        antend = "Todos os dias das 9:00 as 22:00"
                    footer = f"""<p><strong>E-mail:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a></p><p><strong>Hor\u00e1rio de atendimento: </strong>{antend}</p>"""
                    titlezap = f"Fale Conosco \n<span>\nO melhor suporte é com a {empresa}!\n</span>"
                elif pais == "Paquistão" or pais == "Índia":
                    if business_hours == hora1:
                        antend = "Monday to Friday from 9:00 AM to 6:00 PM"
                    elif business_hours == hora2: 
                        antend = "Monday to Friday from 9:00 AM to 10:00 PM"
                    elif business_hours == hora3:
                        antend = "Monday to Saturday from 9:00 AM to 6:00 PM"
                    elif business_hours == hora4:
                        antend = "Monday to Saturday from 9:00 AM to 10:00 PM"
                    elif business_hours == hora5:
                        antend = "Monday to Friday from 9:00 AM to 6:00 PM - Saturday from 9:00 AM to 12:00 PM"
                    elif business_hours == hora6:
                        antend = "Monday to Friday from 9:00 AM to 10:00 PM - Saturday from 9:00 AM to 12:00 PM"
                    elif business_hours == hora7:
                        antend = "Todos os dias das 9:00 as 18:00"
                    elif business_hours == hora8:
                        antend = "Todos os dias das 9:00 as 22:00"
                    hora1 = "Monday to Friday from 9:00 AM to 6:00 PM"
                    hora2 = ""
                    hora3 = "
                    hora4 = ""
                    hora5 = ""
                    hora6 = ""
                    hora7 = ""
                    hora8 = "Every day from 9:00 AM to 10:00 PM"
                    footer = f"""<p><strong>Email:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a></p><p><strong>Business Hours: </strong>{business_hours}</p>"""
                    titlezap = f"Contact Us \n<span>\nThe best support is with {empresa}!\n</span>"
               
                else:
                    footer = f"""<p><strong>Correo electrónico:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a></p><p><strong>Horario de atención: </strong>{business_hours}</p>"""
                    titlezap = f"Contáctanos \n<span>\n¡El mejor soporte es con {empresa}!\n</span>"                                      
                
                
                #######################################################################
                ###########################################################################################
                ####    VARIAVEIS   #########################
                ###########################################################################################
                
                ###########################################################################################
                #### CORES ######
                if paleta == "Paleta de Cor 1":  # LARANJA
                    # barra de anúncio
                    cor1 = "#BB6A09"
                    cor2 = "#FE8B01"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#BB6A09"
                    degrade2 = "#BB6A09"
                    degrade3 = "#BB6A09"
                    degrade4 = "#BB6A09"
                    acent = "#0ABF58"
                    # header
                    accent_color = "#0ABF58"
                    header_background = "#FE8B01"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#BB6A09"
                    degradefooter2 = "#BB6A09"
                    degradefooter3 = "#BB6A09"
                    degradefooter4 = "#BB6A09"
                    footer_background = "#FE8B01"
                    footer_text_color = "#FFFFFF"
                elif paleta == "Paleta de Cor 2": # VERDE
                    # barra de anúncio
                    cor1 = "#60AD3F"
                    cor2 = "#0097B2"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#7ED957"
                    degrade2 = "#0097B2"
                    degrade3 = "#0097B2"
                    degrade4 = "#7ED957"
                    acent = "#0097B2"
                    # header
                    accent_color = "#0097B2"
                    header_background = "#7ED957"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#7ED957"
                    degradefooter2 = "#0097B2"
                    degradefooter3 = "#0097B2"
                    degradefooter4 = "#7ED957"
                    footer_background = "7ED957"
                    footer_text_color = "#FFFFFF"
                elif paleta == "Paleta de Cor 3":  # ROXO
                    # barra de anúncio
                    cor1 = "#860D7E"
                    cor2 = "#600D5A"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#D600FF"
                    degrade2 = "#8610D8"
                    degrade3 = "#D600FF"
                    degrade4 = "#8610D8"
                    acent = "#860D7E"
                    # header
                    accent_color = "#860D7E"
                    header_background = "#600D5A"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#D600FF"
                    degradefooter2 = "#8610D8"
                    degradefooter3 = "#D600FF"
                    degradefooter4 = "#8610D8"
                    footer_background = "600D5A "
                    footer_text_color = "#FFFFFF"
                elif paleta == "Paleta de Cor 4": # PRETO
                    # barra de anúncio
                    cor1 = "#393939"
                    cor2 = "#000000"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#393939"
                    degrade2 = "#393939"
                    degrade3 = "#393939"
                    degrade4 = "#393939"
                    acent = "#393939"
                    # header
                    accent_color = "#393939"
                    header_background = "#000000"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#393939"
                    degradefooter2 = "#393939"
                    degradefooter3 = "#393939"
                    degradefooter4 = "#393939"
                    footer_background = "#000000"
                    footer_text_color = "#FFFFFF"
                elif paleta == "Paleta de Cor 5": # AZUL 
                    # barra de anúncio
                    cor1 = "#5DE0E6"
                    cor2 = "#073D86"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#5DE0E6"
                    degrade2 = "#004AAD"
                    degrade3 = "#5DE0E6"
                    degrade4 = "#004AAD"
                    acent = "#5DE0E6"

                    # header
                    accent_color  = "#5DE0E6"
                    header_background = "#004AAD"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#5DE0E6"
                    degradefooter2 = "#004AAD"
                    degradefooter3 = "#5DE0E6"
                    degradefooter4 = "#004AAD"
                    footer_background = "#004AAD"
                    footer_text_color = "#FFFFFF"
                elif paleta == "Paleta de Cor 6": # VERMELHO
                    # barra de anúncio
                    cor1 = "#B42929"
                    cor2 = "#740606"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#B42929"
                    degrade2 = "#B42929"
                    degrade3 = "#740606"
                    degrade4 = "#740606"
                    acent = "#740606"
                    # header
                    accent_color = "#740606"
                    header_background = "#FF3131"
                    header_text_color = "#F7F7F7"
                    # footer
                    degradefooter1 = "#B42929"
                    degradefooter2 = "#B42929"
                    degradefooter3 = "#740606"
                    degradefooter4 = "#740606"
                    footer_background = "#000000"
                    footer_text_color = "#FFFFFF"
                
                ###########################################################################################
                # Definição das coleções
                colecoes = []
                produtos = None
                if pais == "Brasil":
                    if estilo == "Genérica":
                        primeiro = "Casa"
                        segundo = "Eletrônicos"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Casa", "Eletrônicos", "Pets", "Fitness", "Kids", "Mais Vendidos"]
                    elif estilo == "Eletrônicos":
                        primeiro = "Fones"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Acessórios"
                        quinto = "Produtos-Gamer"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Fones", "Drones", "Smartwatch", "Acessórios", "Produtos Gamer", "Mais Vendidos"]
                    elif estilo == "Kids":
                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Brinquedos"
                        quarto = "Cuidados"
                        quinto = "Maternidade"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Moda", "Acessórios", "Brinquedos", "Cuidados", "Maternidade", "Mais Vendidos"]
                    elif estilo == "Casa":
                        primeiro = "Decoração"
                        segundo = "Utilidades"
                        terceiro = "Jardim"
                        quarto = "Cozinha"
                        quinto = "Seu-Lar"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Decoração", "Utilidades", "Jardim", "Cozinha", "Seu Lar", "Mais Vendidos"]
                    elif estilo == "Pet":
                        primeiro = "Camas"
                        segundo = "Acessórios"
                        terceiro = "Comedouros"
                        quarto = "Brinquedos"
                        quinto = "Roupas"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Camas", "Acessórios", "Comedouros", "Brinquedos", "Roupas", "Mais Vendidos"]
                    elif estilo == "Fitness":
                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Calçados"
                        quarto = "Corretores"
                        quinto = "Produtos-Esportivos"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Moda", "Acessórios", "Calçados", "Corretores", "Produtos Esportivos", "Mais Vendidos"]
                    elif estilo == "Masculino":
                        primeiro = "Acessórios"
                        segundo = "Ferramentas"
                        terceiro = "Pesca"
                        quarto = "Vestuário"
                        quinto = "Automotivos"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Acessórios", "Ferramentas", "Pesca", "Vestuário", "Automotivos", "Mais Vendidos"]
                    elif estilo == "Feminino":
                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Make"
                        quarto = "Escovas-Alisadoras"
                        quinto = "Saúde"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Moda", "Acessórios", "Make", "Escovas Alisadoras", "Saúde", "Mais Vendidos"]
                elif pais == "Portugal":
                    primeiro = "Casa"
                    segundo = "Eletrónica"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Mais-Vendidos"
                    colecoes = ["Casa", "Eletrónica", "Pets", "Fitness", "Kids", "Mais Vendidos"]
                elif pais == "Paquistão" or pais == "Índia":
                    primeiro = "Home"
                    segundo = "Electronics"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Best-Sellers"
                    colecoes = ["Home", "Electronics", "Pets", "Fitness", "Kids", "Best Sellers"]
                else:
                    primeiro = "Casa"
                    segundo = "Electrónica"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Más-vendidos"
                    colecoes = ["Casa", "Electrónica", "Pets", "Fitness", "Kids", "Más vendidos"]
                ##############################################################
                ###############################
                shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
                session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
                shopify.ShopifyResource.activate_session(session)
                ############################# 
                def criar_colecoes(colecoes):
                    for nome in colecoes:
                        nova_colecao = shopify.CustomCollection()
                        nova_colecao.title = nome
                        nova_colecao.save() 
                        print("coleção criada")              
                def adicionar_produtos_colecoes(colecoes, produtos):
                    colecoes = shopify.CustomCollection.find()
                    produtos = shopify.Product.find()
                    headers = {
                        "X-Shopify-Access-Token": PRIVATE_APP_PASSWORD,
                        "Content-Type": "application/json",
                    }
                    mutation = """                                                                                                                                                                                                                                                                                                                             
                                mutation collectionAddProductsV2($id: ID!, $productIds: [ID!]!) {                                                                                                                                                                                                                                                                      
                                    collectionAddProductsV2(id: $id, productIds: $productIds) {                                                                                                                                                                                                                                                                        
                                        job {                                                                                                                                                                                                                                                                                                                          
                                            done                                                                                                                                                                                                                                                                                                                       
                                            id                                                                                                                                                                                                                                                                                                                         
                                        }                                                                                                                                                                                                                                                                                                                              
                                        userErrors {                                                                                                                                                                                                                                                                                                                   
                                            field                                                                                                                                                                                                                                                                                                                      
                                            message                                                                                                                                                                                                                                                                                                                    
                                        }                                                                                                                                                                                                                                                                                                                              
                                    }                                                                                                                                                                                                                                                                                                                                  
                                }                                                                                                                                                                                                                                                                                                                                      
                                """
                    colecoes = shopify.CustomCollection.find()
                    for colecao in colecoes:
                        collection_id = colecao.id
                        collection_handle = colecao.handle

                        # Para todas as coleções, adicione apenas os produtos correspondentes
                        ids_produtos_correspondentes = [produto.id for produto in produtos if collection_handle in produto.tags]

                        for product_id in ids_produtos_correspondentes:
                            try:
                                # Criação das variáveis da mutação
                                variables = {
                                    "id": f"gid://shopify/Collection/{collection_id}",
                                    "productIds": [f"gid://shopify/Product/{product_id}"],
                                }

                                # Criação do corpo da requisição
                                body = {
                                    "query": mutation,
                                    "variables": variables,
                                }
                                # Realização da requisição à API da Shopify
                                response = requests.post(f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",
                                                        headers=headers,
                                                        data=json.dumps(body))

                                # Verificação da resposta
                                # if response.status_code == 200:
                                # print(f"Produto {product_id} adicionado à coleção {collection_id} com sucesso!")
                                # else:
                                # print(f"Erro ao adicionar produto {product_id} à coleção {collection_id}: {response.content}")
                            except Exception as e:
                                print(f"Exceção ao adicionar produto {product_id} à coleção {collection_id}: {str(e)}")
                def editar_tema():
                    temas = shopify.Theme.find()
                    tema_publicado = [tema for tema in temas if tema.role == "main"][0]
                    id_tema_publicado = tema_publicado.id
                    theme = shopify.Theme.find(id_tema_publicado)
                    settings_asset = shopify.Asset.find('config/settings_data.json', theme_id=theme.id)
                    settings_data = json.loads(settings_asset.value)
                    desktop1 = "shopify://shop_images/desktop1.png"
                    desktop2 = "shopify://shop_images/desktop2.png"
                    desktop3 = "shopify://shop_images/desktop3.png"
                    mobile1 = "shopify://shop_images/mobile1.png"
                    mobile2 = "shopify://shop_images/mobile2.png"
                    mobile3 = "shopify://shop_images/mobile3.png"
                    capa1 = "shopify://shop_images/capa1.png"
                    capa2 = "shopify://shop_images/capa2.png"
                    capa3 = "shopify://shop_images/capa3.png"
                    capa4 = "shopify://shop_images/capa4.png"
                    capa5 = "shopify://shop_images/capa5.png"
                    capa6 = "shopify://shop_images/capa6.png"

                    ################################################################################################################
                    ### CORES DO TEMA ###########
                    ################################################################################################################
                    ###### BARRA DE ANUNCIOS #########
                    settings_data['current']['sections']['announcement-bar']['settings']['background1'] = cor1
                    settings_data['current']['sections']['announcement-bar']['settings']['background2'] = cor2
                    settings_data['current']['sections']['announcement-bar']['settings']['text_color'] = textcolor
                    ###### HEADER #########
                    settings_data['current']['header_accent_color'] = acent
                    settings_data['current']['accent_color'] = accent_color
                    settings_data['current']['header_background'] = header_background
                    settings_data['current']['header_text_color'] = header_text_color
                    settings_data['current']['sections']['header']['settings']['background1'] = degrade1
                    settings_data['current']['sections']['header']['settings']['background2'] = degrade2
                    settings_data['current']['sections']['header']['settings']['background3'] = degrade3
                    settings_data['current']['sections']['header']['settings']['background4'] = degrade4
                    ###### FOOTER #########
                    settings_data['current']['sections']['footer']['settings']['background1'] = degradefooter1
                    settings_data['current']['sections']['footer']['settings']['background2'] = degradefooter2
                    settings_data['current']['sections']['footer']['settings']['background3'] = degradefooter3
                    settings_data['current']['sections']['footer']['settings']['background4'] = degradefooter4
                    settings_data['current']['footer_background'] = footer_background
                    settings_data['current']['footer_text_color'] = footer_text_color
                    ##### COLEÇÃO #########
                    settings_data['current']['sections']['1649913264519d35a7']['settings']['background1'] = cor1
                    settings_data['current']['sections']['1649913264519d35a7']['settings']['background2'] = cor2
                    settings_data['current']['sections']['1649913264519d35a7']['settings']['text_color'] = textcolor
                    ################################################################################################################
                    ####    DADOS DA LOJA   ###########
                    ################################################################################################################
                    ### CABEÇALHO ######
                    settings_data['current']['sections']['header']['settings']['navigation_phone_number'] = whatsapp
                    settings_data['current']['sections']['header']['settings']['navigation_email'] = email_suporte
                    #### WHATSAPP #######
                    settings_data['current']['addzap'] = True
                    settings_data['current']['numberzap'] = tel_whats
                    settings_data['current']['titlezap'] = titlezap
                    ##### EDITA FOOTER ######
                    settings_data['current']['sections']['footer']['blocks']['c9a6c378-573f-4c54-9fc6-5ca6e279f3f2']['settings'][
                        'content'] = footer
                    ##### INSERIR BANNERS ##########
                    settings_data['current']['sections']['slideshow']['blocks']['4665d2ed-db3b-479e-8984-d272fdfab8d8']['settings'][
                        'image'] = desktop1
                    settings_data['current']['sections']['slideshow']['blocks']['4665d2ed-db3b-479e-8984-d272fdfab8d8']['settings'][
                        'mobile_image'] = mobile1
                    settings_data['current']['sections']['slideshow']['blocks']['bbea030d-de91-4721-9ea1-9660f647bbd7']['settings'][
                        'image'] = desktop2
                    settings_data['current']['sections']['slideshow']['blocks']['bbea030d-de91-4721-9ea1-9660f647bbd7']['settings'][
                        'mobile_image'] = mobile2
                    settings_data['current']['sections']['slideshow']['blocks']['fa2eac7c-9fbf-4ea9-a172-eca738a9012d']['settings'][
                        'image'] = desktop3
                    settings_data['current']['sections']['slideshow']['blocks']['fa2eac7c-9fbf-4ea9-a172-eca738a9012d']['settings'][
                        'mobile_image'] = mobile3
                    #### COLEÇÔES E CAPAS  ###########
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-1']['settings'][
                        'collection'] = primeiro
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-1']['settings'][
                        'image'] = capa1
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-2']['settings'][
                        'collection'] = segundo
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-2']['settings'][
                        'image'] = capa2
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-3']['settings'][
                        'collection'] = terceiro
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-3']['settings'][
                        'image'] = capa3
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-5']['settings'][
                        'collection'] = quarto
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['1649489657b78ca727-5']['settings'][
                        'image'] = capa4
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['d5a5159b-779d-45d2-871b-4a5e622bb3b2'][
                        'settings']['collection'] = quinto
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['d5a5159b-779d-45d2-871b-4a5e622bb3b2'][
                        'settings']['image'] = capa5
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['0c56a93b-d66f-4823-ae91-bd01bb634a00'][
                        'settings']['collection'] = sexto
                    settings_data['current']['sections']['1649489657b78ca727']['blocks']['0c56a93b-d66f-4823-ae91-bd01bb634a00'][
                        'settings']['image'] = capa6
                    #### FEATURED COLECC #####################
                    settings_data['current']['sections']['16153539675ef431a7']['settings']['collection'] = sexto
                    settings_data['current']['sections']['16153539675ef431a7']['settings']['title'] = "Ofertas da semana"
                    settings_data['current']['sections']['16153541158fc73818']['settings']['collection'] = primeiro
                    settings_data['current']['sections']['16153541158fc73818']['settings']['title'] = primeiro
                    settings_data['current']['sections']['1649913264519d35a7']['settings']['collection'] = segundo
                    settings_data['current']['sections']['1649913264519d35a7']['settings']['title'] = segundo
                    settings_data['current']['sections']['7196b9c9-7ec2-46f0-87be-f84dcc1d14b5']['settings'][
                        'collection'] = terceiro
                    settings_data['current']['sections']['7196b9c9-7ec2-46f0-87be-f84dcc1d14b5']['settings']['title'] = terceiro
                    settings_data['current']['sections']['9dd41389-f9fe-453e-b65e-1cc7383b8c1f']['settings']['collection'] = quarto
                    settings_data['current']['sections']['9dd41389-f9fe-453e-b65e-1cc7383b8c1f']['settings']['title'] = quarto
                    settings_data['current']['sections']['329bc2be-34eb-46d8-8601-39d8eda721d9']['settings']['collection'] = quinto
                    settings_data['current']['sections']['329bc2be-34eb-46d8-8601-39d8eda721d9']['settings']['title'] = quinto

                    settings_asset.value = json.dumps(settings_data)
                    settings_asset.save()
                criar_colecoes(colecoes)
                editar_tema()
                adicionar_produtos_colecoes(colecoes, produtos)
                # Limpar sessão Shopify após a execução bem-sucedida
                shopify.ShopifyResource.clear_session()
                
                return redirect('lojapronta')
              
            except Exception as e:
                print(f'Erro ao executar o programa: {str(e)}')
                    # Retornar uma resposta de erro adequada, se necessário
                return HttpResponse("Erro ao processar o formulário. Por favor, tente novamente.")
        else:
            return redirect('login') 
    
@login_required
def redirecionar_para_produzir(request):
    return redirect('produzir')
############################
                            ### LOJA PRONTA ###### 
############################
@login_required
def lojapronta(request):
    if request.method == 'GET': 
        return render(request, 'lojapronta.html')
    else:
        if request.user.is_authenticated:
            email = request.user.email
            formulario_user = FormularioUser.objects.filter(email=email).first()
            if formulario_user:
                # Marcar a loja como produzida
                formulario_user.lojaproduzida = True
                formulario_user.save()
                
                SHOP_URL = formulario_user.url_loja
                linkloja = f"https://{SHOP_URL}"
                return redirect(linkloja)
            else:
                print("Formulário não encontrado.") 
                return redirect('login')      
@login_required
def redirecionar_para_lojapronta(request):
    return redirect('lojapronta')
############################
    #   MODALS   #
############################
def modal(request):    
    return render(request, 'modal.html')
def modal2(request):    
    return render(request, 'modal2.html')       
def modal3(request):    
    return render(request, 'modal3.html')
def modal4(request):    
    return render(request, 'modal4.html')
def fechar(request):
    if request.method == "POST":
        # Aqui você pode adicionar lógica adicional, se necessário

        return render(request, 'modal4.html')  # Retorna uma resposta renderizada

    # Se não for uma requisição POST, redirecione ou retorne uma resposta adequada
    return render(request, 'modal4.html')  # Redirecionamento de volta para a página
def redirecionar_para_modal4(request):    
    return redirect('modal4')
def modallogin(request):    
    return render(request, 'modallogin.html')
def redirecionar_para_modal(request):
    return redirect('modal')

############################
    #  ###  #
############################
def pagina_erro(request):      
    return render(request, 'erro.html')
def redirecionar_para_pagina_erro(request):
    return redirect('pagina_erro')
def teste(request):    
    return render(request, 'teste.html')
def teste3(request):    
    return render(request, 'email.html') 
@login_required
def retor(request):
    if request.method == 'GET':
        email = request.user.email
        formulario_user = FormularioUser.objects.filter(email=email).first()
        if formulario_user:
            pais = formulario_user.pais
            if pais == "Brasil":
                url = 'https://shopify.pxf.io/6emxyr'
            elif pais == "Portugal":
                url = 'https://shopify.pxf.io/rQmEZj'
            elif pais == "Paquistão" or pais == "Índia":
                url = 'https://shopify.pxf.io/Kj7WEn'
            else:
                url = 'https://shopify.pxf.io/rQmEZj'
            return JsonResponse({'url': url})
        else:
            return JsonResponse({'error': 'Formulário não encontrado'}, status=404)
    else:
        return JsonResponse({'error': 'Método não permitido'}, status=405)
@login_required
def redirecionar_para_retor(request):
    return redirect('retor')
############################
    #  ###  #
############################
